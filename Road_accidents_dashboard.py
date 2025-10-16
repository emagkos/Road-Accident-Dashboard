
import os
from glob import glob
from pathlib import Path
import numpy as np
import pandas as pd
import geopandas as gpd
from shapely.geometry import Point

import streamlit as st
from streamlit_folium import st_folium
import folium
from folium.plugins import MarkerCluster, HeatMap, FastMarkerCluster  # <-- FIX: add FastMarkerCluster

import plotly.express as px
from plotly.colors import sequential, sample_colorscale
from sklearn.cluster import KMeans

# NEW: Word export deps (gracefully optional)
import io
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_OK = True
except Exception:
    _DOCX_OK = False

# Optional: Plotly static export for chart images inside Word
try:
    import kaleido
    _KALEIDO_OK = True
except Exception:
    _KALEIDO_OK = False

# ---------------------------- CONFIG ----------------------------
st.set_page_config(page_title="Lisbon Road Accident Dashboard", page_icon="🚦", layout="wide")

# ------------------------- UTIL FUNCTIONS -----------------------
def find_first_col(candidates, columns):
    cols_lower = {c.lower(): c for c in columns}
    for cand in candidates:
        if cand.lower() in cols_lower:
            return cols_lower[cand.lower()]
    return None

def coerce_datetime(df, col_candidates):
    col = find_first_col(col_candidates, df.columns)
    if col is None:
        return df, None
    df[col] = pd.to_datetime(df[col], errors="coerce", utc=True).dt.tz_convert(None)
    return df, col

def ensure_hour_weekday(df, dt_col):
    if dt_col is not None:
        if "hour" not in df.columns:
            df["hour"] = df[dt_col].dt.hour
        if "weekday" not in df.columns:
            df["weekday"] = df[dt_col].dt.day_name()
    else:
        if "hour" not in df.columns:
            cand = find_first_col(["hour"], df.columns)
            if cand:
                df["hour"] = pd.to_numeric(df[cand], errors="coerce")
        if "weekday" not in df.columns:
            cand = find_first_col(["weekday", "Weekday", "day_of_week"], df.columns)
            if cand:
                df["weekday"] = df[cand].astype(str)
    return df

def _fix_decimals_series(s):
    # Convert "38,7123" -> "38.7123" before to_numeric
    if pd.api.types.is_string_dtype(s) or s.dtype == "object":
        return s.astype("string").str.replace(",", ".", regex=False)
    return s

def to_geodf(df, lat_candidates, lon_candidates):
    # broaden candidates
    lat_candidates = lat_candidates + ["latitude"]
    lon_candidates = lon_candidates + ["longitude"]

    lat_col = find_first_col(lat_candidates, df.columns)
    lon_col = find_first_col(lon_candidates, df.columns)

    if lat_col is None or lon_col is None:
        st.error(
            "Latitude/Longitude columns not found.\n\n"
            f"Looked for LAT in: {lat_candidates}\n"
            f"and LON in: {lon_candidates}\n\n"
            f"Columns in file: {list(df.columns)}"
        )
        st.stop()

    # Handle comma decimals, then numeric
    df[lat_col] = _fix_decimals_series(df[lat_col])
    df[lon_col] = _fix_decimals_series(df[lon_col])
    df[lat_col] = pd.to_numeric(df[lat_col], errors="coerce")
    df[lon_col] = pd.to_numeric(df[lon_col], errors="coerce")

    # Drop invalid rows and keep only plausible ranges
    df = df.dropna(subset=[lat_col, lon_col])
    df = df[(df[lat_col].between(-90, 90)) & (df[lon_col].between(-180, 180))]

    if df.empty:
        st.error("All rows were dropped after cleaning coordinates. Check lat/lon columns and formats.")
    # (continue to allow export on zero-result state)

    gdf = gpd.GeoDataFrame(
        df.copy(),
        geometry=[Point(xy) for xy in zip(df[lon_col], df[lat_col])],
        crs="EPSG:4326",
    )
    return gdf, lat_col, lon_col

def nice_int(x): return f"{int(x):,}"

VALID_WEEKDAYS = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
def normalize_weekday(series):
    s = series.astype("string").fillna("Unknown")
    s = s.where(s.isin(VALID_WEEKDAYS), "Unknown")
    return s

def reds_for(values):
    """Return a list of hex colors (light→dark reds) matching values' normalized rank."""
    v = pd.Series(values, dtype="float")
    if v.nunique() == 1:
        return [sequential.Reds[int(len(sequential.Reds)*0.7)]] * len(v)
    vmin, vmax = v.min(), v.max()
    t = (v - vmin) / (vmax - vmin + 1e-12)
    cols = sample_colorscale("Reds", t.tolist())
    return cols

# --- Severity color helper (consistent red shades) ---
def severity_color(name: str) -> str:
    n = str(name).lower()
    if "fatal" in n or "mortal" in n:
        return "#67000d"  # darkest red
    if "serious" in n or "grave" in n or "severe" in n:
        return "#a50f15"  # dark red
    if "minor" in n or "light" in n:
        return "#f8a883"  # medium red
    if "no injury" in n or "none" in n or "noinjury" in n:
        return "#fcd0bb"  # very light red
    return "#fcbba1"      # fallback light red

# --- Month normalization (fast, specific to "Jan..Dec") ---
MONTHS_ABBR = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
def normalize_month_abbr(series):
    s = series.astype("string").str.strip().str[:3].str.title()
    s = s.where(s.isin(MONTHS_ABBR), pd.NA)
    return pd.Categorical(s, categories=MONTHS_ABBR, ordered=True)

# --- Treat common missing tokens as real NA for text columns (e.g., street) ---
MISSING_TOKENS = {"", "nan", "none", "null", "na", "n/a", "unknown"}
def clean_string_na(s: pd.Series) -> pd.Series:
    s = s.astype("string").str.strip()
    s_lower = s.str.lower()
    return s.mask(s_lower.isin(MISSING_TOKENS))

# ---------------------------- DATA PICKER -----------------------
st.sidebar.markdown("### 📂 Data source")

# 1) Resolve a data folder that works both locally and on Streamlit Cloud
DEFAULT_DATA_DIR = Path(__file__).resolve().parent / "data"
DATA_DIR = Path(os.getenv("DATA_DIR", DEFAULT_DATA_DIR))

# 2) Look for CSVs in the folder (if it exists)
csv_files = []
if DATA_DIR.exists():
    csv_files = sorted([str(p) for p in DATA_DIR.glob("*.csv")])

# 3) Uploader fallback (if nothing found or you want to override)
uploaded = st.sidebar.file_uploader("…or upload a CSV", type=["csv"], help="Upload a CSV with latitude/longitude columns")

DATA_PATH = None
source_label = ""

if uploaded is not None:
    # Use the uploaded file-like object directly with pd.read_csv later
    DATA_PATH = uploaded
    source_label = f"Uploaded file: {uploaded.name}"
elif csv_files:
    # Choose from discovered CSVs
    def_idx = 0
    for i, p in enumerate(csv_files):
        name = Path(p).name.lower()
        if "lisbon" in name or "accident" in name:
            def_idx = i
            break

    csv_choice = st.sidebar.selectbox(
        "Choose CSV file from /data",
        options=csv_files,
        index=def_idx,
        format_func=lambda p: Path(p).name,
    )
    DATA_PATH = csv_choice
    source_label = f"Repo file: {Path(csv_choice).name}"
else:
    # Neither a data folder nor an uploaded file available
    st.error(
        f"No CSVs found.\n\n"
        f"- Expected repo folder: `{DATA_DIR}`\n"
        f"- You can also upload a CSV using the control above."
    )
    st.stop()

st.sidebar.caption(source_label)

# ---------------------------- LOAD DATA -------------------------
@st.cache_data(show_spinner=True)
def load_data(path):
    df = pd.read_csv(path)
    df.columns = df.columns.map(lambda c: c.strip())  # <-- NEW: trim header whitespace

    # Build hour/weekday if a true datetime exists; otherwise keep existing 'hour'/'weekday'
    #df, dt_col = coerce_datetime(df, ["date", "datetime", "timestamp", "accident_date", "data", "data_hora"])
    #df = ensure_hour_weekday(df, dt_col)

    # Detect optional columns used later
    #sev_col = find_first_col(
    #    ["severity", "Severity", "accident_severity", "severidade"],
    #    df.columns
    #)
    street_col = find_first_col(
        ["street"],
        df.columns
    )

    # Robust lat/lon detection + cleaning
    gdf, lat_col, lon_col = to_geodf(
        df,
        lat_candidates=["latitude"],
        lon_candidates=["longitude"]
    )

    # --- Derive severity if dataset has the 30d injury counts ---
    sev_col = None  # default
    if all(c in gdf.columns for c in ["fatal_injuries", "serious_injuries", "minor_injuries"]):
        gdf["severity"] = np.select(
            [
                gdf["fatal_injuries"] > 0,
                gdf["serious_injuries"] > 0,
                gdf["minor_injuries"] > 0,
            ],
            ["Fatal", "Serious", "Minor"],
            default="No Injury",
        )
        sev_col = "severity"  # ensure downstream KPIs use the derived label

    # If the street column exists but is effectively empty (all blank/"nan"/etc.), disable it
    if street_col is not None:
        _non_empty = clean_string_na(gdf[street_col]).dropna()  # <-- NEW
        if _non_empty.empty:
            street_col = None  # hide from UI/plots/popups

    # dt_col is not computed (kept commented above), so keep it as None for consistency
    dt_col = None

    # NOTE: keep your commented line above as requested; we return matching your unpacking below
    return gdf, dt_col, sev_col, street_col, lat_col, lon_col

# Load now (gdf exists from here on)
gdf, DT_COL, SEV_COL, STREET_COL, LAT_COL, LON_COL = load_data(DATA_PATH)

# ---------------------------- SIDEBAR FILTERS -------------------
st.sidebar.title("⚙️ Filters")

# Date range (only if a real datetime column exists)
if DT_COL is not None and gdf[DT_COL].notna().any():
    min_d = gdf[DT_COL].min().date()
    max_d = gdf[DT_COL].max().date()
    date_range = st.sidebar.date_input("Date range", (min_d, max_d), min_value=min_d, max_value=max_d)
else:
    date_range = None

# Hour filter
if "hour" in gdf.columns and gdf["hour"].notna().any():
    hour_min, hour_max = int(gdf["hour"].min()), int(gdf["hour"].max())
    hour_sel = st.sidebar.slider("Hour of day", 0, 23, (hour_min, hour_max))
else:
    hour_sel = (0, 23)

# Weekday filter (robust)
if "weekday" in gdf.columns and gdf["weekday"].notna().any():
    w_all = normalize_weekday(gdf["weekday"])
    weekday_options = VALID_WEEKDAYS + (["Unknown"] if "Unknown" in set(w_all.unique()) else [])
    weekday_sel = st.sidebar.multiselect("Weekdays", weekday_options, default=weekday_options)
else:
    weekday_sel = None

# Severity
if SEV_COL is not None:
    sev_vals = sorted(gdf[SEV_COL].dropna().astype(str).unique())
    severity_sel = st.sidebar.multiselect("Severity", sev_vals, default=sev_vals)
else:
    severity_sel = None

# Map options
st.sidebar.subheader("🗺️ Map options")
map_type = st.sidebar.selectbox("Base map", ["OpenStreetMap", "CartoDB positron", "Stamen Toner", "Stamen Terrain"])
show_markers = st.sidebar.checkbox("Show MarkerCluster", value=True)
show_heatmap = st.sidebar.checkbox("Show HeatMap", value=False)

# Clustering
st.sidebar.subheader("📍 Clustering (KMeans)")
do_cluster = st.sidebar.checkbox("Enable KMeans clustering", value=False)
k_clusters = st.sidebar.slider("Number of clusters (k)", 2, 25, 8)

# OSM overlay
#st.sidebar.subheader("🛰️ OSM streets overlay (experimental)")
#osm_overlay = st.sidebar.checkbox("Load OSM street network (bbox)", value=False)
#osm_msg = st.sidebar.empty()

# ------------------------- APPLY FILTERS ------------------------
df = gdf.copy()

if date_range is not None:
    start_d = pd.to_datetime(date_range[0])
    end_d = pd.to_datetime(date_range[1]) + pd.Timedelta(days=1)
    df = df[(df[DT_COL] >= start_d) & (df[DT_COL] < end_d)]

if "hour" in df.columns:
    df = df[df["hour"].between(hour_sel[0], hour_sel[1])]

if weekday_sel is not None and "weekday" in df.columns:
    _w = normalize_weekday(df["weekday"])
    df = df[_w.isin(weekday_sel)]

if severity_sel is not None and SEV_COL is not None:
    df = df[df[SEV_COL].astype(str).isin(severity_sel)]

# ------------------------------ KPIs ---------------------------
st.title("🚦 Road Accident Dashboard")
#st.caption("Geospatial Data Science for Urban Mobility – Capstone")

col_a, col_b, col_c, col_d = st.columns(4)
with col_a:
    st.metric("Total accidents (filtered)", nice_int(len(df)))
with col_b:
    if SEV_COL:
        severe_like = df[SEV_COL].astype(str).str.contains("serious|grave|fatal|mortal", case=False, na=False)
        st.metric("Serious/Fatal (regex)", nice_int(severe_like.sum()))
    else:
        st.metric("Serious/Fatal (regex)", "—")
with col_c:
    if "hour" in df.columns and not df.empty:
        st.metric("Peak hour (filtered)", int(df["hour"].mode().iloc[0]))
    else:
        st.metric("Peak hour (filtered)", "—")
with col_d:
    # unique (lat, lon) pairs
    uniq_locs = pd.Series(list(zip(df.geometry.y.round(6), df.geometry.x.round(6)))).nunique() if not df.empty else 0
    st.metric("Unique locations", nice_int(uniq_locs))

st.markdown("---")

# ------------------------------ MAP ----------------------------
st.subheader("Interactive Map")
fmap = None
if df.empty:
    st.warning("No data for the selected filters.")
else:
    center = [float(df.geometry.y.median()), float(df.geometry.x.median())]
    # FIX: prefer_canvas speeds up large point rendering
    fmap = folium.Map(location=center, zoom_start=12, control_scale=True, tiles=None, prefer_canvas=True)

    folium.TileLayer(tiles="OpenStreetMap", name="OpenStreetMap").add_to(fmap)
    if map_type == "CartoDB positron":
        folium.TileLayer("CartoDB positron", name="CartoDB positron").add_to(fmap)
    elif map_type == "Stamen Toner":
        folium.TileLayer("Stamen Toner", name="Stamen Toner").add_to(fmap)
    elif map_type == "Stamen Terrain":
        folium.TileLayer("Stamen Terrain", name="Stamen Terrain").add_to(fmap)

    # OSM overlay (optional)
    #if osm_overlay:
    #    try:
    #        import osmnx as ox
    #        minx, miny, maxx, maxy = df.total_bounds
    #        G = ox.graph_from_bbox(maxy, miny, maxx, minx, network_type="drive")
    #        gdf_edges = ox.graph_to_gdfs(G, nodes=False, edges=True)
    #        folium.GeoJson(
    #            gdf_edges.to_crs(4326).__geo_interface__,
    #            name="OSM streets",
    #            style_function=lambda x: {"weight": 1, "color": "#444", "opacity": 0.5}
    #        ).add_to(fmap)
    #        osm_msg.info("OSM street overlay loaded.")
    #    except Exception as e:
    #        osm_msg.error(f"OSM overlay failed: {e}")

    # Markers (robust & fast)
    if show_markers and (not df.empty):
        nrows = len(df)

        popup_cols = []
        for cand in [SEV_COL, (STREET_COL if STREET_COL else None), "hour", "weekday", DT_COL]:
            if isinstance(cand, str) and (cand in df.columns):
                popup_cols.append(cand)

        if nrows > 8000:
            # Use FastMarkerCluster for very large datasets (no per-point popups)
            pts = (
                df.geometry
                  .dropna()
                  .map(lambda g: [float(g.y), float(g.x)])
                  .tolist()
            )
            FastMarkerCluster(pts, name=f"Accidents ({nrows:,})").add_to(fmap)
        else:
            # Normal cluster with small CircleMarkers and HTML popups
            mcluster = MarkerCluster(
                name=f"Accidents ({nrows:,})",
                options={"disableClusteringAtZoom": 17}
            ).add_to(fmap)

            for _, r in df.iterrows():
                lat, lon = float(r.geometry.y), float(r.geometry.x)
                popup_lines = []
                for c in popup_cols:
                    val = r[c]
                    if pd.isna(val):
                        continue
                    if c == DT_COL:
                        try:
                            val = pd.to_datetime(val).strftime("%Y-%m-%d %H:%M")
                        except Exception:
                            pass
                    popup_lines.append(f"<b>{c}</b>: {val}")
                popup_html = "<br>".join(popup_lines) if popup_lines else "Accident"

                folium.CircleMarker(
                    location=[lat, lon],
                    radius=3,
                    color="#b2182b",  # dark red marker
                    fill=True,
                    fill_opacity=0.6,
                    popup=folium.Popup(popup_html, max_width=300),
                ).add_to(mcluster)

    # Heatmap (red gradient)  -- FIX: string keys + keep inside the map block
    if show_heatmap and (not df.empty):
        heat_pts = (
            df.geometry
              .dropna()
              .map(lambda g: [float(g.y), float(g.x)])
              .tolist()
        )
        red_gradient = {
            "0.0": "#fff5f0",
            "0.5": "#fb6a4a",
            "1.0": "#99000d",
        }
        HeatMap(
            heat_pts,
            radius=12,
            blur=18,
            name="HeatMap",
            gradient=red_gradient,
        ).add_to(fmap)

    # Clustering (centroids)
    if do_cluster and len(df) >= k_clusters:
        coords = np.column_stack([df.geometry.y.values, df.geometry.x.values])
        km = KMeans(n_clusters=k_clusters, n_init=10, random_state=42)
        labels = km.fit_predict(coords)
        centers = km.cluster_centers_
        for i, (clat, clon) in enumerate(centers):
            folium.Marker(
                location=[float(clat), float(clon)],
                icon=folium.Icon(color="red", icon="info-sign"),
                tooltip=f"Cluster {i}"
            ).add_to(fmap)
    elif do_cluster:
        st.info("Not enough records to cluster with current filters.")

    # FIX: add layer control once, after all layers
    folium.LayerControl(collapsed=False).add_to(fmap)

    st_folium(fmap, width=None, height=600)

# ------------------------------ CHARTS --------------------------
st.markdown("### 📊 Visual summaries")
col1, col2 = st.columns(2)

# keep figures for Word export
fig_hour = None
fig_weekday = None
fig_severity = None
fig_streets = None
fig_month = None  # NEW

with col1:
    if "hour" in df.columns and not df.empty:
        tmp = df.groupby("hour").size().reset_index(name="count").sort_values("hour")
        fig_hour = px.bar(
            tmp, x="hour", y="count",
            title="Accidents by Hour of Day",
            labels={"hour": "Hour", "count": "Accidents"},
        )
        fig_hour.update_traces(marker_color=reds_for(tmp["count"]))
        fig_hour.update_layout(coloraxis_showscale=False)
        st.plotly_chart(fig_hour, use_container_width=True)
    else:
        st.info("Hour column not available.")

with col2:
    if "weekday" in df.columns and not df.empty:
        s = normalize_weekday(df["weekday"])
        order = VALID_WEEKDAYS + (["Unknown"] if "Unknown" in set(s.unique()) else [])
        cat = pd.Categorical(s, categories=order, ordered=True)
        tmp = pd.DataFrame({"weekday": cat}).groupby("weekday").size().reset_index(name="count")
        fig_weekday = px.bar(
            tmp, x="weekday", y="count",
            title="Accidents by Weekday",
            labels={"weekday": "Weekday", "count": "Accidents"},
        )
        fig_weekday.update_traces(marker_color=reds_for(tmp["count"]))
        fig_weekday.update_layout(xaxis_tickangle=-15, coloraxis_showscale=False)
        st.plotly_chart(fig_weekday, use_container_width=True)
    else:
        st.info("Weekday column not available.")

col3, col4 = st.columns(2)
with col3:
    if SEV_COL is not None and not df.empty:
        tmp = df[SEV_COL].astype(str).fillna("Unknown").value_counts().reset_index()
        tmp.columns = ["severity", "count"]
        colors = [severity_color(s) for s in tmp["severity"]]
        fig_severity = px.pie(
            tmp, names="severity", values="count",
            title="Accidents by Severity", hole=0.35
        )
        fig_severity.update_traces(marker=dict(colors=colors))
        st.plotly_chart(fig_severity, use_container_width=True)
    else:
        st.info("Severity column not available.")

with col4:
    if STREET_COL is not None and not df.empty:
        s = clean_string_na(df[STREET_COL]).dropna()  # <-- NEW
        topn = s.value_counts().head(10).reset_index()  # <-- NEW
        topn.columns = ["street", "accidents"]
        fig_streets = px.bar(
            topn, x="street", y="accidents",
            title="Top 10 Streets (by count)",
        )
        fig_streets.update_traces(marker_color=reds_for(topn["accidents"]))
        fig_streets.update_layout(xaxis_tickangle=-35, coloraxis_showscale=False)
        st.plotly_chart(fig_streets, use_container_width=True)
    else:
        st.info("Street/road column not available.")

# ---------- FIXED: Accident type % by Month with custom severity order ----------
st.markdown("### 📅 Accident type share by month")
if SEV_COL is not None and not df.empty:
    mcol = find_first_col(["month"], df.columns)
    if mcol is not None and df[mcol].notna().any():
        MONTHS_ABBR = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
        month_cat = df[mcol].astype("string").str.strip().str[:3].str.title()
        month_cat = pd.Categorical(month_cat.where(month_cat.isin(MONTHS_ABBR), pd.NA),
                                   categories=MONTHS_ABBR, ordered=True)

        tmp = (
            pd.DataFrame({"_month": month_cat, "_sev": df[SEV_COL].astype(str)})
            .dropna(subset=["_month"])
            .groupby(["_month", "_sev"])
            .size()
            .reset_index(name="count")
        )

        if not tmp.empty:
            # % by month
            month_totals = tmp.groupby("_month")["count"].transform("sum")
            tmp["percent"] = (tmp["count"] / month_totals) * 100.0

            # Example order (light→dark). Change if you like.
            SEV_ORDER = ["No Injury", "Minor", "Serious", "Fatal"]

            present_sev = tmp["_sev"].unique().tolist()
            sev_order = [s for s in SEV_ORDER if s in present_sev] + [s for s in present_sev if s not in SEV_ORDER]

            tmp["_sev"] = pd.Categorical(tmp["_sev"], categories=sev_order, ordered=True)
            tmp = tmp.sort_values(["_month", "_sev"])

            color_map = {lab: severity_color(lab) for lab in sev_order}

            fig_month = px.bar(
                tmp,
                x="_month", y="percent",
                color="_sev",
                category_orders={"_month": MONTHS_ABBR, "_sev": sev_order},
                color_discrete_map=color_map,
                title="Accident type (%) by month",
                labels={"_month": "Month", "percent": "Share (%)", "_sev": "Severity"},
            )
            fig_month.update_layout(barmode="stack", yaxis_tickformat=".0f%", legend_traceorder="normal")
            st.plotly_chart(fig_month, use_container_width=True)
        else:
            st.info("Month column present but empty after cleaning.")
    else:
        st.info("Month column not available.")
else:
    st.info("Severity column not available; cannot compute monthly severity shares.")

# --------------------------- INSIGHTS ---------------------------
st.markdown("---")
st.subheader("📝 Quick insights")
insights = []
if df.empty:
    st.write("No insights available (no rows after filtering).")
else:
    if "hour" in df.columns and df["hour"].notna().any():
        peak_hour = int(df["hour"].mode().iloc[0])
        insights.append(f"• Peak accident hour: {peak_hour}:00.")
    if "weekday" in df.columns and df["weekday"].notna().any():
        s = normalize_weekday(df["weekday"])
        risky_day = s.mode().iloc[0]
        insights.append(f"• Most frequent weekday: {risky_day}.")
    if SEV_COL is not None:
        sev_share = (df[SEV_COL].astype(str)
                     .str.contains("serious|grave|fatal|mortal", case=False, na=False).mean())
        insights.append(f"• Serious/Fatal share (regex): {sev_share:.1%} of filtered accidents.")
    st.write("\n".join(insights) if insights else "No obvious patterns detected with current filters.")

# --------------------------- DOWNLOAD (data only) ---------------------------
st.markdown("### ⬇️ Download filtered data")
csv_bytes = df.drop(columns="geometry", errors="ignore").to_csv(index=False).encode("utf-8")
st.download_button(label="Download CSV (filtered)", data=csv_bytes, file_name="lisbon_accidents_filtered.csv", mime="text/csv")

# --------------------------- WORD REPORT (single .docx) ---------------------------
st.markdown("---")
st.subheader("📝 Export Word report (.docx)")

def _fig_to_png_bytes(fig, width=1000, height=600, scale=1.0):
    if not _KALEIDO_OK or fig is None:
        return None
    try:
        return fig.to_image(format="png", width=width, height=height, scale=scale)
    except Exception:
        return None

def build_word_report():
    if not _DOCX_OK:
        st.error("python-docx is not installed. Please: pip install python-docx")
        return None

    doc = Document()

    # Title
    title = doc.add_heading("Lisbon Road Accident Dashboard – Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Source & filters
    doc.add_paragraph(f"Data source: {source_label}")
    filt_lines = []
    if 'date_range' in globals() and date_range is not None:
        filt_lines.append(f"Date range: {date_range[0]} to {date_range[1]}")
    if isinstance(hour_sel, tuple):
        filt_lines.append(f"Hour range: {hour_sel[0]}–{hour_sel[1]}")
    if 'weekday_sel' in globals() and weekday_sel is not None:
        filt_lines.append(f"Weekdays: {', '.join(weekday_sel) if weekday_sel else '—'}")
    if 'severity_sel' in globals() and severity_sel is not None:
        filt_lines.append(f"Severity: {', '.join(severity_sel) if severity_sel else '—'}")
    doc.add_paragraph("Filters applied:\n" + ("\n".join(f"• {x}" for x in filt_lines) if filt_lines else "• None"))

    # KPIs
    doc.add_heading("KPIs", level=1)
    doc.add_paragraph(f"Total accidents (filtered): {nice_int(len(df))}")
    if SEV_COL:
        severe_like = df[SEV_COL].astype(str).str.contains("serious|grave|fatal|mortal", case=False, na=False)
        doc.add_paragraph(f"Serious/Fatal (regex): {nice_int(severe_like.sum())}")
    if "hour" in df.columns and not df.empty:
        doc.add_paragraph(f"Peak hour (filtered): {int(df['hour'].mode().iloc[0])}")
    uniq_locs = pd.Series(list(zip(df.geometry.y.round(6), df.geometry.x.round(6)))).nunique() if not df.empty else 0
    doc.add_paragraph(f"Unique locations: {nice_int(uniq_locs)}")

    # Insights
    doc.add_heading("Quick insights", level=1)
    if insights:
        for line in insights:
            doc.add_paragraph(line.replace("• ", ""), style="List Bullet")
    else:
        doc.add_paragraph("No insights for current filters.")

    # Severity breakdown
    doc.add_heading("Severity breakdown", level=1)
    if SEV_COL is not None and not df.empty:
        sev_tbl = df[SEV_COL].astype(str).fillna("Unknown").value_counts().reset_index()
        sev_tbl.columns = ["Severity", "Count"]
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Severity", "Count"
        for _, r in sev_tbl.iterrows():
            row = table.add_row().cells
            row[0].text = str(r["Severity"])
            row[1].text = str(int(r["Count"]))
    else:
        doc.add_paragraph("Severity column not available.")

    # Top streets table (auto-hidden if STREET_COL was blank)
    if STREET_COL is not None and not df.empty:
        doc.add_heading("Top 10 streets (by count)", level=1)
        s = clean_string_na(df[STREET_COL]).dropna()  # <-- NEW
        topn = s.value_counts().head(10).reset_index()  # <-- NEW
        topn.columns = ["Street", "Accidents"]
        table2 = doc.add_table(rows=1, cols=2)
        hdr2 = table2.rows[0].cells
        hdr2[0].text, hdr2[1].text = "Street", "Accidents"
        for _, r in topn.iterrows():
            row = table2.add_row().cells
            row[0].text = str(r["Street"])
            row[1].text = str(int(r["Accidents"]))

    # Charts (as images if kaleido available)
    doc.add_heading("Charts", level=1)
    if not _KALEIDO_OK:
        doc.add_paragraph("Chart images omitted (Kaleido not installed). Install with: pip install -U kaleido")
    else:
        for name, fig in [
            ("Accidents by Hour of Day", fig_hour),
            ("Accidents by Weekday", fig_weekday),
            ("Accidents by Severity", fig_severity),
            ("Top 10 Streets (by count)", fig_streets),
            ("Accident type (%) by month", fig_month),  # NEW in report
        ]:
            if fig is None:
                continue
            img = _fig_to_png_bytes(fig)
            if img:
                doc.add_paragraph(name)
                doc.add_picture(io.BytesIO(img), width=Inches(6.5))
            else:
                doc.add_paragraph(f"{name}: (image export failed)")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

colL, colR = st.columns([1,1])
with colL:
    if st.button("Generate Word report"):
        data = build_word_report()
        if data is not None:
            st.download_button(
                "Download lisbon_accidents_report.docx",
                data=data,
                file_name="lisbon_accidents_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
with colR:
    if not _DOCX_OK:
        st.info("To enable Word export, install:  pip install python-docx")
    if not _KALEIDO_OK:
        st.info("To embed chart images in the Word report, install:  pip install -U kaleido")
